"""
Build the PCG Monthly Report Word template (.docx) for KPV Papamoa.

Run:
    /tmp/pcg-build-venv/bin/python build-pcg-template.py

Output: PCG Monthly Report - Papamoa - TEMPLATE.docx in the same folder.

Palette and structure follow:
    plans/Stage 1 Papamoa/pcg-report-redesign-powerpoint-template.md
    reference/brand-guidelines.md
"""

from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt, Cm, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).parent / "PCG Monthly Report - Papamoa - TEMPLATE.docx"
LOGO = "/tmp/kpv-logo-cropped.png"

# ---- KPV palette (hex strings used by OOXML shading) ----
KPV_NAVY = "1B3A5B"
DASH_NAVY = "131F39"
LEAF_GREEN = "7FC242"
INK = "3A3A3A"
PAPER = "F5F5F5"
RULE = "E5E5E5"
SUBTITLE_GREY = "A8A49B"
WHITE = "FFFFFF"
RAG_GREEN = "5BA844"
RAG_AMBER = "E8A33C"
RAG_RED = "C0392B"
RAG_GREY = "A8A49B"

# Fonts (with safe fallbacks)
HEADING_FONT = "Garamond"     # falls back fine on systems without Source Serif Pro
BODY_FONT = "Calibri"


# ---------- low-level helpers ----------

def hex_to_rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell, sz="4", color="E5E5E5", sides=("top", "left", "bottom", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Set cell padding in twentieths of a point (e.g. 80 = 4pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        n = OxmlElement(f"w:{name}")
        n.set(qn("w:w"), str(val))
        n.set(qn("w:type"), "dxa")
        tcMar.append(n)
    tcPr.append(tcMar)


def run(p, text, *, font=BODY_FONT, size=11, bold=False, italic=False, color_hex=INK):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = hex_to_rgb(color_hex)
    return r


def para(doc_or_cell, text="", *, align=None, space_after=Pt(4), **run_kw):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    if text:
        run(p, text, **run_kw)
    return p


def placeholder(doc_or_cell, text, *, size=10):
    p = para(doc_or_cell, "")
    run(p, text, font=BODY_FONT, size=size, bold=True, color_hex=KPV_NAVY)
    return p


def section_banner(doc, title, rag=None):
    """Full-width navy bar with white heading text + optional RAG chip at right."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    page_w = Mm(170)
    tbl.columns[0].width = Mm(150)
    tbl.columns[1].width = Mm(20)
    remove_table_borders(tbl)

    title_cell = tbl.cell(0, 0)
    shade_cell(title_cell, KPV_NAVY)
    set_cell_margins(title_cell, top=120, bottom=120, left=200, right=120)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run(p, title.upper(), font=HEADING_FONT, size=16, bold=True, color_hex=WHITE)

    rag_cell = tbl.cell(0, 1)
    shade_cell(rag_cell, KPV_NAVY)
    set_cell_margins(rag_cell, top=80, bottom=80, left=40, right=40)
    rag_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = rag_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    if rag:
        rag_colors = {"G": RAG_GREEN, "A": RAG_AMBER, "R": RAG_RED, "?": RAG_GREY}
        # Embed a small inner table for the RAG chip
        inner = rag_cell.add_table(rows=1, cols=1)
        inner.autofit = False
        inner.columns[0].width = Mm(14)
        remove_table_borders(inner)
        chip = inner.cell(0, 0)
        shade_cell(chip, rag_colors.get(rag, RAG_GREY))
        chip.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(chip, top=40, bottom=40, left=40, right=40)
        cp = chip.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(cp, rag, font=BODY_FONT, size=12, bold=True, color_hex=WHITE)
    else:
        run(p, "", color_hex=WHITE)


def kpi_tile_row(doc, items, *, fill=PAPER, label_color=SUBTITLE_GREY, value_color=KPV_NAVY, height_mm=20):
    """Row of KPI tiles. items is list of (label, value) or (label, value, color_hex)."""
    tbl = doc.add_table(rows=1, cols=len(items))
    tbl.autofit = False
    col_w = Mm(170 // len(items))
    for c in tbl.columns:
        c.width = col_w
    remove_table_borders(tbl)
    for idx, item in enumerate(items):
        if len(item) == 3:
            label, value, vc = item
        else:
            label, value = item
            vc = value_color
        cell = tbl.cell(0, idx)
        shade_cell(cell, fill)
        set_cell_borders(cell, sz="4", color=RULE)
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Label
        lp = cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_after = Pt(2)
        run(lp, label.upper(), font=BODY_FONT, size=8, bold=True, color_hex=label_color)
        # Value
        vp = cell.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.paragraph_format.space_after = Pt(0)
        run(vp, value, font=HEADING_FONT, size=18, bold=True, color_hex=vc)


def rag_strip(doc, items):
    """Row of RAG tiles. items is list of (label, rag_letter)."""
    tbl = doc.add_table(rows=1, cols=len(items))
    tbl.autofit = False
    col_w = Mm(170 // len(items))
    for c in tbl.columns:
        c.width = col_w
    remove_table_borders(tbl)
    rag_colors = {"G": RAG_GREEN, "A": RAG_AMBER, "R": RAG_RED, "?": RAG_GREY}
    for idx, (label, rag) in enumerate(items):
        cell = tbl.cell(0, idx)
        shade_cell(cell, rag_colors.get(rag, RAG_GREY))
        set_cell_margins(cell, top=120, bottom=120, left=80, right=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_after = Pt(0)
        run(lp, label.upper(), font=BODY_FONT, size=9, bold=True, color_hex=WHITE)
        vp = cell.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.paragraph_format.space_after = Pt(0)
        run(vp, rag, font=HEADING_FONT, size=14, bold=True, color_hex=WHITE)


def chart_placeholder(doc, label, *, height_mm=70):
    """Bordered card with a 'paste chart here' note."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Mm(170)
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    shade_cell(cell, PAPER)
    set_cell_borders(cell, sz="6", color=RULE)
    set_cell_margins(cell, top=400, bottom=400, left=200, right=200)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run(p, label, font=BODY_FONT, size=10, bold=True, italic=True, color_hex=SUBTITLE_GREY)


def two_col_lists(doc, left_title, left_items, right_title, right_items):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Mm(85)
    tbl.columns[1].width = Mm(85)
    remove_table_borders(tbl)
    for cell, title, items in ((tbl.cell(0, 0), left_title, left_items),
                                (tbl.cell(0, 1), right_title, right_items)):
        shade_cell(cell, PAPER)
        set_cell_borders(cell, sz="4", color=RULE)
        set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
        # Title
        tp = cell.paragraphs[0]
        tp.paragraph_format.space_after = Pt(4)
        run(tp, title.upper(), font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
        for item in items:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run(p, "• ", font=BODY_FONT, size=10, bold=True, color_hex=KPV_NAVY)
            run(p, item, font=BODY_FONT, size=10, color_hex=INK)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def setup_page(doc):
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)


def set_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    # Three-cell table for left/center/right footer
    tbl = footer.add_table(rows=1, cols=3, width=Mm(170))
    tbl.autofit = False
    tbl.columns[0].width = Mm(60)
    tbl.columns[1].width = Mm(60)
    tbl.columns[2].width = Mm(50)
    remove_table_borders(tbl)
    left, mid, right = tbl.rows[0].cells
    p1 = left.paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
    run(p1, "Karaka Pines Villages — Papamoa", font=BODY_FONT, size=8, color_hex=SUBTITLE_GREY)
    p2 = mid.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(0)
    run(p2, "PCG Monthly Report — ", font=BODY_FONT, size=8, color_hex=SUBTITLE_GREY)
    run(p2, "[Month YYYY]", font=BODY_FONT, size=8, bold=True, color_hex=KPV_NAVY)
    p3 = right.paragraphs[0]; p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p3.paragraph_format.space_after = Pt(0)
    run(p3, "Page ", font=BODY_FONT, size=8, color_hex=SUBTITLE_GREY)
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    r = p3.add_run(); r.font.size = Pt(8); r.font.color.rgb = hex_to_rgb(SUBTITLE_GREY)
    r._r.append(fldChar1); r._r.append(instrText); r._r.append(fldChar2)


def set_header(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    try:
        r.add_picture(LOGO, height=Mm(14))
    except Exception:
        run(p, "[KPV LOGO]", font=BODY_FONT, size=10, bold=True, color_hex=KPV_NAVY)


# ---------- the actual document ----------

def build():
    doc = Document()
    setup_page(doc)
    set_header(doc)
    set_footer(doc)

    # Default style tweaks
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_to_rgb(INK)

    # ============================================================
    # PAGE 1 — Cover + Executive RAG
    # ============================================================
    para(doc, "", space_after=Pt(2))  # small top spacer below header

    # Big cover title block (navy hero)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Mm(170)
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    shade_cell(cell, DASH_NAVY)
    set_cell_margins(cell, top=400, bottom=400, left=400, right=400)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    run(p, "PCG MONTHLY REPORT", font=HEADING_FONT, size=28, bold=True, color_hex=WHITE)
    p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    run(p2, "PAPAMOA — STAGE 1", font=BODY_FONT, size=12, color_hex=SUBTITLE_GREY)
    p3 = cell.add_paragraph(); p3.paragraph_format.space_after = Pt(0)
    run(p3, "[KYLE TO UPDATE: Month YYYY]", font=BODY_FONT, size=14, bold=True, color_hex=LEAF_GREEN)

    para(doc, "", space_after=Pt(6))

    # Overall RAG header
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "OVERALL PROJECT STATUS", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)

    # 6 RAG tiles
    rag_strip(doc, [
        ("Quality", "G"),
        ("Programme", "G"),
        ("Cost", "A"),
        ("H&S", "G"),
        ("Sales", "A"),
        ("Māori", "G"),
    ])
    para(doc, "", space_after=Pt(2))
    placeholder(doc, "[KYLE TO SET: RAG status for each category — replace G / A / R letters above]")

    para(doc, "", space_after=Pt(6))

    # Headline numbers
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "HEADLINE NUMBERS", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    kpi_tile_row(doc, [
        ("Units Settled", "6 / 128"),
        ("Construction Complete", "11"),
        ("Avg List Price", "$914k"),
        ("Near Miss (period)", "2"),
    ])
    para(doc, "", space_after=Pt(2))
    placeholder(doc, "[KYLE TO UPDATE: pull current month values from 107 Dashboard widgets]")

    para(doc, "", space_after=Pt(8))

    # Key events
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "KEY EVENTS THIS PERIOD", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    for _ in range(3):
        placeholder(doc, "•  [KYLE TO ADD: max 15 words]")

    add_page_break(doc)

    # ============================================================
    # PAGE 2 — Programme
    # ============================================================
    section_banner(doc, "Programme", rag="G")
    para(doc, "", space_after=Pt(4))

    chart_placeholder(doc,
        "[PASTE: Construction % Complete by Block — screenshot Chart 3 from the 107 Dashboard "
        "(see plans/Stage 1 Papamoa/pcg-report-dashboard-chart-buildspec.md)]",
        height_mm=70)
    para(doc, "", space_after=Pt(6))

    two_col_lists(doc,
        "Achieved this period",
        [
            "[KYLE TO ADD: bullet, max 15 words]",
            "[KYLE TO ADD: bullet, max 15 words]",
            "[KYLE TO ADD: bullet, max 15 words]",
        ],
        "Planned next period",
        [
            "[KYLE TO ADD: bullet, max 15 words]",
            "[KYLE TO ADD: bullet, max 15 words]",
            "[KYLE TO ADD: bullet, max 15 words]",
        ],
    )

    add_page_break(doc)

    # ============================================================
    # PAGE 3 — Cost / Financial
    # ============================================================
    section_banner(doc, "Cost / Financial", rag="A")
    para(doc, "", space_after=Pt(4))

    # Budget vs Certified table — placeholder values
    tbl = doc.add_table(rows=4, cols=4)
    tbl.autofit = False
    widths_mm = [55, 38, 38, 39]
    for i, w in enumerate(widths_mm):
        tbl.columns[i].width = Mm(w)
    remove_table_borders(tbl)
    header = ["", "Budget", "Certified to Date", "To Complete"]
    for i, h in enumerate(header):
        cell = tbl.cell(0, i)
        shade_cell(cell, KPV_NAVY)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        if i > 0: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run(p, h.upper(), font=BODY_FONT, size=9, bold=True, color_hex=WHITE)
    rows = [
        ("Stage 1 Civil", "$2,225k", "$1,932k", "$0"),
        ("Stage 2 Civil", "$1,304k", "$391k", "$1,233k"),
        ("[KYLE TO ADD: row]", "[KYLE]", "[KYLE]", "[KYLE]"),
    ]
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            shade_cell(cell, PAPER if ri % 2 == 0 else WHITE)
            set_cell_borders(cell, sz="4", color=RULE)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            if ci > 0: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            is_placeholder = "[KYLE" in val
            run(p, val,
                font=BODY_FONT, size=10,
                bold=is_placeholder,
                color_hex=KPV_NAVY if is_placeholder else INK)

    para(doc, "", space_after=Pt(8))

    # BBDNZ drawdown paste-in block
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "BBDNZ DRAWDOWN (SCHEDULE 2)", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    chart_placeholder(doc,
        "[KYLE TO PASTE: BBDNZ monthly drawdown table — sourced from BBD attachment]",
        height_mm=45)

    para(doc, "", space_after=Pt(6))
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "HEADLINE COMMENTARY", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    placeholder(doc, "[KYLE TO WRITE: max 50 words on cost position — variances, contingency draws, approvals]")

    add_page_break(doc)

    # ============================================================
    # PAGE 4 — Sales & Marketing
    # ============================================================
    section_banner(doc, "Sales & Marketing", rag="A")
    para(doc, "", space_after=Pt(4))

    # Two charts side by side
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Mm(85)
    tbl.columns[1].width = Mm(85)
    remove_table_borders(tbl)
    for col_idx, label in enumerate([
        "[PASTE: Sales Pipeline donut — Chart 1, from 107 Dashboard once built]",
        "[PASTE: Cumulative settlements vs target — Chart 2 (requires Monthly Sales Snapshot sheet)]",
    ]):
        cell = tbl.cell(0, col_idx)
        shade_cell(cell, PAPER)
        set_cell_borders(cell, sz="6", color=RULE)
        set_cell_margins(cell, top=400, bottom=400, left=200, right=200)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run(p, label, font=BODY_FONT, size=9, bold=True, italic=True, color_hex=SUBTITLE_GREY)

    para(doc, "", space_after=Pt(8))

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "SALES COMMENTARY", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    placeholder(doc, "[KYLE TO WRITE: max 80 words. Cover: enquiry levels, appointments/conversions this period, "
                     "market conditions one-liner. Replace the ~500-word Schedule 3 narrative.]")

    add_page_break(doc)

    # ============================================================
    # PAGE 5 — Sales detail (per-unit)
    # ============================================================
    section_banner(doc, "Sales — Per-Unit Detail")
    para(doc, "", space_after=Pt(4))

    chart_placeholder(doc,
        "[PASTE: full screenshot of 107 - Sales Report - Expanded (Smartsheet report id 1830541908201348). "
        "Filter to Stage 1 only if Stage 2A creates noise.]",
        height_mm=160)

    add_page_break(doc)

    # ============================================================
    # PAGE 6 — Health & Safety
    # ============================================================
    section_banner(doc, "Health and Safety", rag="G")
    para(doc, "", space_after=Pt(4))

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "LEAD INDICATORS", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    kpi_tile_row(doc, [
        ("Inductions", "17"),
        ("Contractor Audits", "7"),
        ("Internal Audits (PM)", "0"),
        ("External Inspections", "0"),
    ])
    para(doc, "", space_after=Pt(6))

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "LAG INDICATORS", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    kpi_tile_row(doc, [
        ("Near Miss Minor", "2"),
        ("Near Miss Serious (PSIF)", "0"),
        ("MTI", "0"),
        ("LTI", "0"),
    ])
    para(doc, "", space_after=Pt(2))
    placeholder(doc, "[KYLE TO UPDATE: pull current period values from 107 H&S Monthly Indicators sheet]")

    para(doc, "", space_after=Pt(8))
    chart_placeholder(doc,
        "[PASTE: 6-month H&S trend stacked column — Chart 4]",
        height_mm=55)

    para(doc, "", space_after=Pt(4))
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "OPEN ITEMS", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    placeholder(doc, "[KYLE TO WRITE: 1 line summary — see 107 - H&S Open Items Report for detail]")

    add_page_break(doc)

    # ============================================================
    # PAGE 7 — Risk Register
    # ============================================================
    section_banner(doc, "Risk Register")
    para(doc, "", space_after=Pt(4))

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "TOP 5 ACTIVE RISKS", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)

    # Risk table — populated with current data from 107 Risk Report (Yellow RAG, May 2026)
    risks = [
        ("RSK-0002", "H&S audit frequency below target", "A", "PM / Signature"),
        ("RSK-0004", "Trip hazard — open drains", "A", "Signature"),
        ("RSK-0005", "Wind-blown materials hazard", "A", "Signature"),
        ("RSK-0003", "Near-miss reporting gap (PCG vs Sig)", "A", "PM"),
        ("RSK-0001", "HSMS does not deliver safe behaviour", "A", "PM"),
    ]
    tbl = doc.add_table(rows=len(risks) + 1, cols=4)
    tbl.autofit = False
    widths_mm = [25, 80, 15, 50]
    for i, w in enumerate(widths_mm):
        tbl.columns[i].width = Mm(w)
    remove_table_borders(tbl)
    headers = ["ID", "Risk", "RAG", "Owner"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        shade_cell(cell, KPV_NAVY)
        set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        if h == "RAG": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, h.upper(), font=BODY_FONT, size=9, bold=True, color_hex=WHITE)
    rag_colors = {"G": RAG_GREEN, "A": RAG_AMBER, "R": RAG_RED}
    for ri, (rid, risk, rag, owner) in enumerate(risks, start=1):
        bg = PAPER if ri % 2 == 0 else WHITE
        for ci, val in enumerate([rid, risk, "", owner]):
            cell = tbl.cell(ri, ci)
            shade_cell(cell, bg)
            set_cell_borders(cell, sz="4", color=RULE)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            run(p, val, font=BODY_FONT, size=10, color_hex=INK)
        # RAG chip in col 2
        rag_cell = tbl.cell(ri, 2)
        shade_cell(rag_cell, rag_colors.get(rag, RAG_GREY))
        set_cell_borders(rag_cell, sz="4", color=RULE)
        rp = rag_cell.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.paragraph_format.space_after = Pt(0)
        run(rp, rag, font=BODY_FONT, size=10, bold=True, color_hex=WHITE)

    para(doc, "", space_after=Pt(2))
    placeholder(doc, "[KYLE TO REFRESH: pull current top 5 from 107 Risk Report (id 2047711023943556) "
                     "ordered by Score DESC, then filter to Active]")

    para(doc, "", space_after=Pt(8))
    chart_placeholder(doc,
        "[PASTE: Risk heat map (5×5 likelihood × consequence) — Chart 5]",
        height_mm=65)

    add_page_break(doc)

    # ============================================================
    # PAGE 8 — Māori Procurement
    # ============================================================
    section_banner(doc, "Māori Procurement")
    para(doc, "", space_after=Pt(4))

    para(doc, "", space_after=Pt(2))
    placeholder(doc, "[NEXT: build a 107 - Māori Procurement Register Smartsheet, then auto-populate this table]")
    para(doc, "", space_after=Pt(4))

    tbl = doc.add_table(rows=4, cols=4)
    tbl.autofit = False
    widths_mm = [55, 45, 30, 40]
    for i, w in enumerate(widths_mm):
        tbl.columns[i].width = Mm(w)
    remove_table_borders(tbl)
    headers = ["Contractor", "Service", "Contract Value", "Date"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        shade_cell(cell, KPV_NAVY)
        set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        run(p, h.upper(), font=BODY_FONT, size=9, bold=True, color_hex=WHITE)
    for ri in range(1, 4):
        bg = PAPER if ri % 2 == 0 else WHITE
        for ci in range(4):
            cell = tbl.cell(ri, ci)
            shade_cell(cell, bg)
            set_cell_borders(cell, sz="4", color=RULE)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            label = "[KYLE TO ADD]" if ri == 1 else ""
            if ri == 1 and ci == 0:
                run(p, "Out the Gate Contractors", font=BODY_FONT, size=10, color_hex=INK)
            elif ri == 1 and ci == 1:
                run(p, "Boundary Fence", font=BODY_FONT, size=10, color_hex=INK)
            elif ri == 1 and ci == 3:
                run(p, "12/02/2026", font=BODY_FONT, size=10, color_hex=INK)
            else:
                run(p, label, font=BODY_FONT, size=10, bold=bool(label), color_hex=KPV_NAVY if label else INK)

    para(doc, "", space_after=Pt(8))
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "CULTURAL / IWI UPDATE", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    placeholder(doc, "[KYLE TO WRITE: max 40 words. e.g. road naming with Waitaha, wetland planting coordination.]")

    add_page_break(doc)

    # ============================================================
    # PAGE 9 — Appendix A: Live Data Links
    # ============================================================
    section_banner(doc, "Appendix A — Live Data")
    para(doc, "", space_after=Pt(6))

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    run(p, "Smartsheet sources backing this report", font=HEADING_FONT, size=12, bold=True, color_hex=KPV_NAVY)

    links = [
        ("107 Dashboard", "smartsheet.com/dashboards/MgV3W8XJrMJHRrv4gmpmJ3cj3QvHQVggGFM75g71"),
        ("107 - Civil Programme", "Sheet id 6233992763232132"),
        ("107 - Construction Programme", "Sheet id 5063906232848260"),
        ("107 - Unit Register", "Sheet id 5289542339743620"),
        ("107 - Sales Tracking", "Sheet id 8224519314427780"),
        ("107 - H&S Monthly Indicators", "Sheet id 3104570405244804"),
        ("107 - Sales Report - Expanded", "Report id 1830541908201348"),
        ("107 - Papamoa - Signature Report", "Report id 6144240361885572"),
        ("107 - Risk Report", "Report id 2047711023943556"),
        ("107 - H&S Open Items Report", "Report id 8821624189964164"),
    ]
    tbl = doc.add_table(rows=len(links), cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Mm(70)
    tbl.columns[1].width = Mm(100)
    remove_table_borders(tbl)
    for i, (name, value) in enumerate(links):
        bg = PAPER if i % 2 == 0 else WHITE
        for ci, val in enumerate([name, value]):
            cell = tbl.cell(i, ci)
            shade_cell(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            bold = ci == 0
            run(p, val, font=BODY_FONT, size=10, bold=bold, color_hex=INK)

    para(doc, "", space_after=Pt(10))
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "APPROVALS / RESOLUTIONS LOG", font=BODY_FONT, size=9, bold=True, color_hex=SUBTITLE_GREY)
    placeholder(doc, "[KYLE TO MAINTAIN: dated list of board / CEO approvals and resolutions]")

    add_page_break(doc)

    # ============================================================
    # PAGE 10+ — Appendix B: Photo Record
    # ============================================================
    section_banner(doc, "Appendix B — Photographic Record")
    para(doc, "", space_after=Pt(4))
    placeholder(doc, "[KYLE TO ADD: site photos for the period, organised by block. "
                     "Captioned with date + block. Source: attachments on 107 Construction Programme rows.]")
    para(doc, "", space_after=Pt(8))

    for block in ["Berkeley 105-107", "Monterey 109-111", "Melrose 113-115", "Street View"]:
        # Block label
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        run(p, block.upper(), font=BODY_FONT, size=10, bold=True, color_hex=KPV_NAVY)
        # Photo grid placeholder — 2x2
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        tbl.columns[0].width = Mm(85)
        tbl.columns[1].width = Mm(85)
        remove_table_borders(tbl)
        for ci in range(2):
            cell = tbl.cell(0, ci)
            shade_cell(cell, PAPER)
            set_cell_borders(cell, sz="6", color=RULE)
            set_cell_margins(cell, top=600, bottom=600, left=200, right=200)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(0)
            run(cp, "[PHOTO]", font=BODY_FONT, size=10, bold=True, italic=True, color_hex=SUBTITLE_GREY)
        para(doc, "", space_after=Pt(6))

    # ============================================================

    doc.save(OUTPUT)
    print(f"Built: {OUTPUT}")


if __name__ == "__main__":
    build()
